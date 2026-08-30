"""Browse resumes that live in a shared Google Drive folder.

Two stages, deliberately separate:

  1. Read and parse. Walk the folder, stream each file into memory, and turn
     it into a JSON record -- metadata plus plain text. No LLM is involved, so
     this is fast, free, and never rate limited. Records are cached per file,
     keyed on the Drive modifiedTime, so an unchanged resume is downloaded and
     parsed exactly once.

  2. Display. Only when a candidate's details are actually shown does the LLM
     run over that document's cached text to produce the structured fields.
     Those results are cached the same way, so the model runs once per resume
     version rather than once per page load.

Stateless with respect to storage: nothing is written to uploads/, and nothing
is written to the database. Both caches live in process memory and are cleared
by a restart.
"""

import io
import logging
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from typing import Any, Dict, Iterator, List, Optional, Tuple

import fitz  # pymupdf, already used for uploaded resumes

from app.config import settings
from app.services import NotFound, Unavailable
from app.services import llm_service, pdf_service

logger = logging.getLogger(__name__)

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
WRITE_SCOPES = ["https://www.googleapis.com/auth/drive"]

FOLDER_MIME = "application/vnd.google-apps.folder"
PDF_MIME = "application/pdf"
GDOC_MIME = "application/vnd.google-apps.document"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"

# Legacy .doc is walked so it can be reported as skipped rather than vanishing.
DOCUMENT_MIMES = frozenset({PDF_MIME, GDOC_MIME, DOCX_MIME, DOC_MIME})

# Below this, a PDF has no text layer -- it's a scan and needs OCR.
MIN_CHARS = 100

_local = threading.local()

# Guards the caches below. Held only for short reads and writes -- never across
# a Drive fetch or an LLM call, so /status stays responsive during a refresh.
_cache_lock = threading.Lock()

# Whole-listing cache, so repeated page loads don't even re-walk the folder.
_listing: Dict[str, Any] = {"fetched_at": 0.0, "payload": None}

# Stage 1: parsed documents, keyed on file_id. Holds the modifiedTime the
# record was built from, so a changed file is re-parsed and nothing else is.
_text_cache: Dict[str, Dict[str, Any]] = {}

# Stage 2: LLM extractions, keyed the same way. Filled lazily on display.
_detail_cache: Dict[str, Dict[str, Any]] = {}

# Serialises refreshes so two simultaneous requests don't both walk the folder.
_refresh_lock = threading.Lock()


# --------------------------------------------------------------------------
# Drive client
# --------------------------------------------------------------------------
def _client():
    """One client per thread.

    googleapiclient's underlying http object is not thread-safe, and the parse
    stage runs across a thread pool.
    """
    client = getattr(_local, "drive", None)
    if client is not None:
        return client

    if not settings.drive_enabled:
        raise Unavailable(
            "Google Drive is not configured. Set DRIVE_ROOT_FOLDER_ID and "
            "DRIVE_SERVICE_ACCOUNT_FILE in .env, then restart the server."
        )

    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise Unavailable(
            "The Google Drive libraries are not installed. Run "
            "`pip install -r requirements.txt`."
        ) from exc

    try:
        creds = service_account.Credentials.from_service_account_file(
            settings.drive_service_account_file, scopes=SCOPES
        )
    except FileNotFoundError as exc:
        raise Unavailable(
            f"Service account key not found at "
            f"{settings.drive_service_account_file}."
        ) from exc
    except ValueError as exc:
        raise Unavailable(f"Service account key is not valid JSON: {exc}") from exc

    client = build("drive", "v3", credentials=creds, cache_discovery=False)
    _local.drive = client
    return client


def _list_children(drive, folder_id: str) -> Iterator[Dict[str, Any]]:
    """Yield every direct child of a folder, following pagination."""
    token = None
    while True:
        response = (
            drive.files()
            .list(
                q=f"'{folder_id}' in parents and trashed=false",
                fields="nextPageToken, files(id, name, mimeType, modifiedTime, webViewLink)",
                pageSize=1000,
                pageToken=token,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
            )
            .execute()
        )
        for item in response.get("files", []):
            yield item

        token = response.get("nextPageToken")
        if not token:
            return


def _walk(drive, folder_id: str, label: str = "") -> Iterator[Tuple[Dict[str, Any], str]]:
    """Recursively yield (file metadata, folder label) for every document.

    Folder names become the label, matching the resume_parser convention where
    the tree encodes a domain -- manager/, developer/, and so on.
    """
    for item in _list_children(drive, folder_id):
        if item["mimeType"] == FOLDER_MIME:
            child = f"{label}/{item['name']}".lstrip("/")
            for pair in _walk(drive, item["id"], child):
                yield pair
        elif item["mimeType"] in DOCUMENT_MIMES:
            yield item, label


# --------------------------------------------------------------------------
# Fetch + parse, entirely in memory
# --------------------------------------------------------------------------
def _drain(request) -> io.BytesIO:
    """Run a media request to completion into a BytesIO. Never touches disk."""
    from googleapiclient.http import MediaIoBaseDownload

    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    return buf


def _fetch(drive, meta: Dict[str, Any]) -> Tuple[io.BytesIO, str]:
    """Return (bytes, kind) for a supported document.

    Google Docs carry no binary content -- get_media() rejects them with
    403 fileNotDownloadable -- so they are exported to PDF instead.
    """
    mime = meta["mimeType"]
    file_id = meta["id"]

    if mime == PDF_MIME:
        return _drain(drive.files().get_media(fileId=file_id)), "pdf"
    if mime == GDOC_MIME:
        return _drain(drive.files().export_media(fileId=file_id, mimeType=PDF_MIME)), "pdf"
    if mime == DOCX_MIME:
        return _drain(drive.files().get_media(fileId=file_id)), "docx"

    raise ValueError(f"unsupported mimeType: {mime}")


def _text_from_pdf(buf: io.BytesIO) -> str:
    """Parse from the byte stream, so no temporary file is ever created."""
    with fitz.open(stream=buf.getvalue(), filetype="pdf") as doc:
        if doc.needs_pass:
            raise ValueError("password protected")
        return "\n".join(page.get_text("text") for page in doc).strip()


def _text_from_docx(buf: io.BytesIO) -> str:
    import docx

    document = docx.Document(buf)
    parts = [p.text for p in document.paragraphs]
    # Resumes are often laid out in tables, whose text is not in .paragraphs.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts).strip()


def _parse_document(meta: Dict[str, Any], domain: str) -> Dict[str, Any]:
    """Stage 1: turn one Drive file into a JSON record. No LLM.

    One unreadable file must not abandon the whole listing, so every failure is
    captured on the record rather than raised.
    """
    record = {
        "file_id": meta["id"],
        "filename": meta["name"],
        "domain": domain,
        "mime_type": meta["mimeType"],
        "modified": meta.get("modifiedTime"),
        "web_view_link": meta.get("webViewLink"),
        "state": "skipped",
        "reason": None,
        "chars": 0,
        "text": "",
        "from_cache": False,
    }

    if meta["mimeType"] == DOC_MIME:
        record["reason"] = "Legacy .doc format is not supported."
        return record

    try:
        drive = _client()
        buf, kind = _fetch(drive, meta)
        text = _text_from_pdf(buf) if kind == "pdf" else _text_from_docx(buf)
    except Exception as exc:
        logger.warning("Drive: failed to read %s: %s", meta["name"], exc)
        record["state"] = "failed"
        record["reason"] = f"Could not read this file: {exc}"
        return record

    if kind == "pdf" and len(text) < MIN_CHARS:
        record["reason"] = "No text layer -- this looks like a scan and needs OCR."
        return record

    record["state"] = "parsed"
    record["text"] = text
    record["chars"] = len(text)
    # A name is cheap to guess from the first lines and needs no model. It gives
    # the list something to show before the LLM details arrive.
    record["display_name"] = pdf_service.guess_name(text) or meta["name"]
    return record


def _public(record: Dict[str, Any]) -> Dict[str, Any]:
    """A record without its text -- resumes are large and the list doesn't
    need the body, only the details endpoint does."""
    return {k: v for k, v in record.items() if k != "text"}


# --------------------------------------------------------------------------
# Stage 1: the document listing
# --------------------------------------------------------------------------
def _build_listing() -> Dict[str, Any]:
    """Walk the folder and parse only what changed.

    Listing is metadata only and costs one Drive call per folder. A file whose
    modifiedTime matches what the cache was built from is never downloaded or
    parsed again.
    """
    drive = _client()
    documents = list(_walk(drive, settings.drive_root_folder_id.strip()))

    records: List[Dict[str, Any]] = []
    todo: List[Tuple[Dict[str, Any], str]] = []
    seen = set()

    with _cache_lock:
        for meta, domain in documents:
            seen.add(meta["id"])
            entry = _text_cache.get(meta["id"])
            if entry is not None and entry["modified"] == meta.get("modifiedTime"):
                # Reuse the parsed text, but take metadata fresh -- a move
                # between folders changes the domain label.
                record = dict(entry["record"])
                record.update(
                    {
                        "filename": meta["name"],
                        "domain": domain,
                        "modified": meta.get("modifiedTime"),
                        "web_view_link": meta.get("webViewLink"),
                        "from_cache": True,
                    }
                )
                records.append(record)
            else:
                todo.append((meta, domain))

        # Drop files that have left the folder, from both caches.
        for stale in [fid for fid in _text_cache if fid not in seen]:
            del _text_cache[stale]
        for stale in [fid for fid in _detail_cache if fid not in seen]:
            del _detail_cache[stale]

    if todo:
        workers = max(1, min(settings.drive_max_workers, len(todo)))
        with ThreadPoolExecutor(max_workers=workers) as pool:
            fresh = list(pool.map(lambda pair: _parse_document(*pair), todo))

        with _cache_lock:
            for record in fresh:
                # A failure is usually transient (rate limit, timeout), so it
                # is not cached -- the next read retries it. "skipped" is a
                # property of the file itself and is safe to remember.
                if record["state"] in ("parsed", "skipped"):
                    _text_cache[record["file_id"]] = {
                        "modified": record["modified"],
                        "record": record,
                    }
                # A re-parsed file has new text, so any detail built from the
                # old version is stale.
                _detail_cache.pop(record["file_id"], None)
        records.extend(fresh)

    # Parsed first, then by name -- failures and scans sink to the bottom.
    order = {"parsed": 0, "skipped": 1, "failed": 2}
    records.sort(key=lambda r: (order.get(r["state"], 3), (r["filename"] or "").lower()))

    counts = {
        "parsed": 0,
        "skipped": 0,
        "failed": 0,
        "total": len(records),
        "reused": sum(1 for r in records if r["from_cache"]),
        "reparsed": len(todo),
    }
    for record in records:
        counts[record["state"]] = counts.get(record["state"], 0) + 1

    with _cache_lock:
        counts["detailed"] = sum(1 for r in records if r["file_id"] in _detail_cache)

    return {
        "documents": [_public(r) for r in records],
        "counts": counts,
        "fetched_at": time.time(),
    }


def load_documents(force: bool = False, full: bool = False) -> Dict[str, Any]:
    """Stage 1 listing. Never calls the LLM.

    force -- re-walk the folder instead of serving the whole-listing cache.
             Unchanged files are still reused, so this is cheap.
    full  -- also discard the per-file caches, re-parsing every resume.
    """
    if not settings.drive_enabled:
        raise Unavailable(
            "Google Drive is not configured. Set DRIVE_ROOT_FOLDER_ID and "
            "DRIVE_SERVICE_ACCOUNT_FILE in .env, then restart the server."
        )

    def _fresh_enough() -> Optional[Dict[str, Any]]:
        with _cache_lock:
            payload = _listing["payload"]
            age = time.time() - _listing["fetched_at"]
        if payload is not None and age < settings.drive_cache_ttl_seconds:
            return {**payload, "cached": True, "age_seconds": int(age)}
        return None

    if not force and not full:
        hit = _fresh_enough()
        if hit is not None:
            return hit

    with _refresh_lock:
        # Another request may have refreshed while we waited for the lock.
        if not force and not full:
            hit = _fresh_enough()
            if hit is not None:
                return hit

        if full:
            with _cache_lock:
                _text_cache.clear()
                _detail_cache.clear()

        payload = _build_listing()

        with _cache_lock:
            _listing["payload"] = payload
            _listing["fetched_at"] = payload["fetched_at"]

    return {**payload, "cached": False, "age_seconds": 0}


# --------------------------------------------------------------------------
# Stage 2: LLM details, on display
# --------------------------------------------------------------------------
def get_details(file_id: str, force: bool = False) -> Dict[str, Any]:
    """Structured candidate fields for one document.

    This is the only place the LLM is used. It runs against text already in the
    cache, so no Drive call is made, and the result is cached against the same
    modifiedTime -- the model runs once per resume version, not once per view.
    """
    with _cache_lock:
        entry = _text_cache.get(file_id)
        cached = _detail_cache.get(file_id)

    if entry is None:
        # Cold process, or a listing that has never been built. Build it and
        # look again rather than failing on a link the user just clicked.
        load_documents()
        with _cache_lock:
            entry = _text_cache.get(file_id)
            cached = _detail_cache.get(file_id)

    if entry is None:
        raise NotFound(f"No document {file_id} in the Drive folder.")

    record = entry["record"]
    if record["state"] != "parsed":
        raise NotFound(
            f"'{record['filename']}' has no text to read: {record['reason']}"
        )

    if not force and cached is not None and cached["modified"] == record["modified"]:
        return {**cached["detail"], "from_cache": True}

    extracted, engine = llm_service.extract_resume(record["text"])
    detail = {
        "file_id": file_id,
        "filename": record["filename"],
        "domain": record["domain"],
        "full_name": extracted.full_name or record.get("display_name"),
        "email": str(extracted.email) if extracted.email else None,
        "phone": extracted.phone,
        "location": extracted.location,
        "years_experience": extracted.years_experience,
        "skills": extracted.skills,
        "education": extracted.education,
        "summary": extracted.summary,
        "extracted_by": engine,
        "from_cache": False,
    }

    with _cache_lock:
        _detail_cache[file_id] = {"modified": record["modified"], "detail": detail}

    return detail


# --------------------------------------------------------------------------
# Uploading
#
# Reading uses the service account; writing cannot. A service account owns no
# Drive storage quota, so creating a file in a personal Drive fails with
# "Service Accounts do not have storage quota" even when the folder is shared
# with it as Editor. Writes therefore run as a real Google account, authorised
# once via scripts/drive_authorize.py.
# --------------------------------------------------------------------------
def _write_client():
    client = getattr(_local, "drive_write", None)
    if client is not None:
        return client

    if not settings.drive_enabled:
        raise Unavailable(
            "Google Drive is not configured. Set DRIVE_ROOT_FOLDER_ID and "
            "DRIVE_SERVICE_ACCOUNT_FILE in .env, then restart the server."
        )
    if not settings.drive_upload_enabled:
        raise Unavailable(
            "Uploading to Drive is not set up. Service accounts have no storage "
            "quota, so uploads have to run as a real Google account: run "
            "`python scripts/drive_authorize.py`, then set DRIVE_OAUTH_TOKEN_FILE "
            "in .env and restart."
        )

    try:
        from google.auth.transport.requests import Request
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise Unavailable(
            "The Google Drive libraries are not installed. Run "
            "`pip install -r requirements.txt`."
        ) from exc

    try:
        creds = Credentials.from_authorized_user_file(
            settings.drive_oauth_token_file, WRITE_SCOPES
        )
    except FileNotFoundError as exc:
        raise Unavailable(
            f"Drive upload token not found at {settings.drive_oauth_token_file}. "
            "Run `python scripts/drive_authorize.py` to create it."
        ) from exc
    except ValueError as exc:
        raise Unavailable(f"Drive upload token is not valid: {exc}") from exc

    if not creds.valid:
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
            Path(settings.drive_oauth_token_file).write_text(creds.to_json())
        else:
            raise Unavailable(
                "The Drive upload token has expired and cannot be refreshed. "
                "Run `python scripts/drive_authorize.py` again."
            )

    client = build("drive", "v3", credentials=creds, cache_discovery=False)
    _local.drive_write = client
    return client


def upload_resume(filename: str, content: bytes) -> Dict[str, Any]:
    """Put one PDF into the shared folder and return its Drive metadata.

    The new file is added to the same folder the listing reads, so an uploaded
    resume shows up on the Resumes tab after the next refresh.
    """
    from googleapiclient.http import MediaIoBaseUpload

    drive = _write_client()
    created = (
        drive.files()
        .create(
            body={
                "name": filename,
                "parents": [settings.drive_root_folder_id.strip()],
            },
            media_body=MediaIoBaseUpload(
                io.BytesIO(content), mimetype="application/pdf", resumable=False
            ),
            fields="id, name, webViewLink, size, modifiedTime",
            supportsAllDrives=True,
        )
        .execute()
    )

    # The folder has changed, so the cached listing is stale.
    with _cache_lock:
        _listing["payload"] = None
        _listing["fetched_at"] = 0.0

    logger.info("Drive: uploaded %s as %s", filename, created["id"])
    return {
        "file_id": created["id"],
        "filename": created["name"],
        "web_view_link": created.get("webViewLink"),
        "modified": created.get("modifiedTime"),
    }


def status() -> Dict[str, Any]:
    """Configuration and cache state, without touching Drive or the LLM."""
    with _cache_lock:
        payload = _listing["payload"]
        age = time.time() - _listing["fetched_at"] if payload else None
        files_cached = len(_text_cache)
        details_cached = len(_detail_cache)

    return {
        "configured": settings.drive_enabled,
        "folder_id_set": bool(settings.drive_root_folder_id.strip()),
        "key_file_set": bool(settings.drive_service_account_file.strip()),
        "cache_ttl_seconds": settings.drive_cache_ttl_seconds,
        "cached": payload is not None,
        "age_seconds": int(age) if age is not None else None,
        "counts": payload["counts"] if payload else None,
        "files_cached": files_cached,
        "details_cached": details_cached,
        "upload_enabled": settings.drive_upload_enabled,
    }
